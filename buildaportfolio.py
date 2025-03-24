import copy
import os
import tkinter as tk
from datetime import date as date_type
from datetime import datetime, time, timedelta
from tkinter import Frame, Label, StringVar, Tk, filedialog, messagebox, ttk

import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import numpy as np
import openpyxl
import pandas as pd
from docx import Document
from docx.shared import Inches
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.dates import date2num, num2date
from matplotlib.figure import Figure
from matplotlib.widgets import Cursor
from scipy.stats import norm
from sklearn.linear_model import LinearRegression
from tkcalendar import DateEntry

pd.set_option('display.float_format','{:.10f}'.format)
script_dir = os.path.dirname(os.path.abspath(__file__))
index_to_table_df_nse = os.path.join(script_dir, 'index_to_table_mapping_nse.csv')
try:
    index_to_table_df_nse = pd.read_csv(index_to_table_df_nse)
    temp_dict = index_to_table_df_nse.to_dict()
    abc = dict()
    for i in range(len(temp_dict['index'])):
        abc[temp_dict["index"][i]] = temp_dict["table_name"][i]
except (FileNotFoundError, pd.errors.EmptyDataError):
    print(f"Error reading file: {index_to_table_df_nse}")
    abc = {}
index_to_table_df_bse=os.path.join(script_dir,'index_to_table_mapping_bse.csv') 
try:
    index_to_table_df_bse = pd.read_csv(index_to_table_df_bse)
    temp_dict = index_to_table_df_bse.to_dict()
    b = dict()
    for i in range(len(temp_dict['index'])):
        b[temp_dict["index"][i]] = temp_dict["table_name"][i]
except (FileNotFoundError, pd.errors.EmptyDataError):
    print(f"Error reading file: {index_to_table_df_bse}")
    b = {}
abc.update(b)

INDEX_TO_TABLES_MAPPING=abc
STOCK_INDEXES = [] #append from BSE indices checkbutton
GROUP = [] #provide option & append here checkbutton
global RATE_RISK_FREE_RETURN_NSE
INTERVAL=0
# STOCK_EXCHANGE=""
current_annotation = None

class BuildPortfolioMenu(tk.Menu):
    def __init__(self, master):
        super().__init__(master, tearoff=0)
        
        style = ttk.Style(self)
        style.configure('TLabel', font=('Segoe UI', 12))  # Times New Roman font with size 14
        style.configure('TButton', font=('Segoe UI', 12))
        style.configure('TEntry', font=('Segoe UI', 12))
        style.configure('TFrame', font=('Segoe UI', 12))
        style.configure('TLabel', font=('Segoe UI', 12))
        style.configure('TCheckbutton', font=('Segoe UI', 12))
        style.configure('Treeview', font=('Segoe UI', 12))
        style.configure('Treeview.Heading', font=('Segoe UI', 12, 'bold'))
        
        self.add_command(label="Initial Inputs", command=self.initial_inputs,accelerator="Ctrl+i")
        self.bind_all("<Control-i>",lambda event:self.initial_inputs())
        self.add_command(label="Enter Risk Free Rate of Return", command=self.risk_free_return)
        self.add_command(label="Data Grid View", command=self.grid_view)
        
        self.START_DATE=None
        self.END_DATE=None
        self.conn=None
        self.cursor=None
        
    def update_credentials(self, conn, cursor):
        self.conn=conn
        self.cursor=cursor
        self.DATABASE_NAME=self.conn.database
        
    def initial_inputs(self):
        new_window=tk.Toplevel(self)
        new_window.title("initial_inputs")
        new_window.geometry("1200x800")
        
        font_setting = ('Segoe UI', 12)
        
        exchanges_frame = ttk.LabelFrame(new_window, text="Exchanges")
        exchanges_frame.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.BOTH, expand=True)


        group_frame=ttk.LabelFrame(new_window,text="Group")
        group_frame.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.BOTH, expand=True)
        
        bse_indices_frame = ttk.LabelFrame(new_window, text="BSE Indices")
        bse_indices_frame.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.BOTH, expand=True)

        nse_indices_frame = ttk.LabelFrame(new_window, text="NSE Indices")
        nse_indices_frame.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.BOTH, expand=True)
        self.canvas = tk.Canvas(nse_indices_frame)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Add a scrollbar to the canvas
        self.scrollbar = ttk.Scrollbar(nse_indices_frame, orient=tk.VERTICAL, command=self.canvas.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Configure the canvas to use the scrollbar
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.bind('<Configure>', lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        # Create a frame inside the canvas to hold the checkboxes
        self.nse_indices_inner_frame = ttk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.nse_indices_inner_frame, anchor="nw")

        time_period_frame = ttk.LabelFrame(new_window, text="Select Time Period")
        time_period_frame.pack(side=tk.LEFT, padx=10, pady=10, fill=tk.BOTH, expand=True)

        self.exchange_var = tk.StringVar()
        self.exchange_var.trace("w", self.update_indices_state)
        tk.Radiobutton(exchanges_frame, text="Bombay Stock Exchange", variable=self.exchange_var,value="bse", font=font_setting).pack(anchor='w', padx=5, pady=2)
        tk.Radiobutton(exchanges_frame, text="National Stock Exchange", variable=self.exchange_var,value="nse", font=font_setting).pack(anchor='w', padx=5, pady=2)


        #add widgets to the group 
        self.group_vars = [tk.BooleanVar() for _ in range(2)]
        group_option=['A','B']
        self.group_widgets = []
        for i, text in enumerate(group_option):
            var = self.group_vars[i]
            widget = ttk.Checkbutton(group_frame, text=text, variable=var)
            widget.pack(anchor='w', padx=5, pady=2)
            self.group_widgets.append(widget)
        
        # Add widgets to the BSE indices frame
        bse_indices =[
            "bse_greenex",
            "bse_sensex",
            "bse_100",
            "bse_200",
            "bse_500",
            "bse_information_technology",
            "bse_fast_moving_consumer_goods",
            "bse_capital_goods",
            "bse_consumer_durables",
            "bse_healthcare",
            "bse_dollex_200",
            "bse_teck",
            "bse_psu",
            "bse_dollex_30",
            "bse_bankex",
            "bse_auto",
            "bse_metal",
            "bse_oil_&_gas",
            "bse_midcap",
            "bse_smallcap",
            "bse_dollex_100",
            "bse_realty",
            "bse_power",
            "bse_ipo",
            "bse_carbonex",
            "bse_sme_ipo",
            "bse_india_infrastructure_index",
            "bse_cpse",
            "bse_india_manufacturing_index",
            "bse_allcap",
            "bse_largecap",
            "bse_sensex_50",
            "bse_sensex_next_50",
            "bse_bharat_22_index",
            "bse_150_midcap_index",
            "bse_250_smallcap_index",
            "bse_250_largemidcap_index",
            "bse_400_midsmallcap_index",
            "bse_smallcap_select_index",
            "bse_midcap_select_index",
            "bse_100_largecap_tmc_index",
            "bse_consumer_descretionary",
            "bse_energy",
            "bse_financial_services",
            "bse_industrials",
            "bse_telecommunication",
            "bse_utilities",
            "bse_diversified_financials_revenue_growth_index",
            "bse_private_banks_index",
            "bse_momentum_index",
            "bse_low_volatility_index",
            "bse_quality_index",
            "bse_enhanced_value_index",
            "bse_dividend_stability_index",
            "bse_100_esg_index"
        ]
        self.bse_all_var = tk.BooleanVar()
        self.bse_all_widget=ttk.Checkbutton(bse_indices_frame, text="BSE_ALL", variable=self.bse_all_var, command=lambda: self.handle_all_indices(self.bse_all_var, self.bse_indices_vars, self.bse_indices_widgets))
        self.bse_all_widget.pack(anchor='w', padx=5, pady=2)
        self.bse_indices_vars = [tk.BooleanVar() for _ in range(len(bse_indices))]

        self.bse_indices_widgets = []
        for i, text in enumerate(bse_indices):
            var = self.bse_indices_vars[i]
            widget = ttk.Checkbutton(bse_indices_frame, text=text, variable=var)
            widget.pack(anchor='w', padx=5, pady=2)
            self.bse_indices_widgets.append(widget)
                
        # STOCK_INDEXES.extend(self.bse_indices_widgets)
        

        # Add widgets to the NSE indices frame
        self.nse_indices_widgets = []
        nse_indices = [
            "nifty_50",
            "nifty_next_50",
            "nifty_100",
            "nifty_200",
            "nifty_500",
            "nifty_midcap_50",
            "nifty_midcap_100",
            "nifty_smallcap_100",
            "india_vix",
            "nifty_midcap_150",
            "nifty_smallcap_50",
            "nifty_smallcap_250",
            "nifty_midsmallcap_400",
            "nifty500_multicap_50:25:25",
            "nifty_largemidcap_250",
            "nifty_midcap_select",
            "nifty_total_market",
            "nifty_microcap_250",
            "nifty_bank",
            "nifty_auto",
            "nifty_financial_services",
            "nifty_financial_services_25/50",
            "nifty_fmcg",
            "nifty_it",
            "nifty_media",
            "nifty_metal",
            "nifty_pharma",
            "nifty_psu_bank",
            "nifty_private_bank",
            "nifty_realty",
            "nifty_healthcare_index",
            "nifty_consumer_durables",
            "nifty_oil_&_gas",
            "nifty_midsmall_healthcare",
            "nifty_dividend_opportunities_50",
            "nifty_growth_sectors_15",
            "nifty100_quality_30",
            "nifty50_value_20",
            "nifty50_tr_2x_leverage",
            "nifty50_pr_2x_leverage",
            "nifty50_tr_1x_inverse",
            "nifty50_pr_1x_inverse",
            "nifty50_dividend_points",
            "nifty_alpha_50",
            "nifty50_equal_weight",
            "nifty100_equal_weight",
            "nifty100_low_volatility_30",
            "nifty200_quality_30",
            "nifty_alpha_low_volatility_30",
            "nifty200_momentum_30",
            "nifty_midcap150_quality_50",
            "nifty200_alpha_30",
            "nifty_midcap150_momentum_50",
            "nifty_commodities",
            "nifty_india_consumption",
            "nifty_cpse",
            "nifty_energy",
            "nifty_infrastructure",
            "nifty100_liquid_15",
            "nifty_midcap_liquid_15",
            "nifty_mnc",
            "nifty_pse",
            "nifty_services_sector",
            "nifty100_esg_sector_leaders",
            "nifty_india_digital",
            "nifty100_esg",
            "nifty_india_manufacturing",
            "nifty_india_corporate_group_index___tata_group_25%_cap",
            "nifty500_multicap_india_manufacturing_50:30:20",
            "nifty500_multicap_infrastructure_50:30:20",
            "nifty_8_13_yr_g_sec",
            "nifty_10_yr_benchmark_g_sec",
            "nifty_10_yr_benchmark_g_sec_(clean_price)",
            "nifty_4_8_yr_g_sec_index",
            "nifty_11_15_yr_g_sec_index",
            "nifty_15_yr_and_above_g_sec_index",
            "nifty_composite_g_sec_index"
        ]
        self.nse_all_var = tk.BooleanVar()
        self.nse_all_widget=ttk.Checkbutton(self.nse_indices_inner_frame, text="NSE_ALL", variable=self.nse_all_var, command=lambda: self.handle_all_indices(self.nse_all_var, self.nse_indices_vars, self.nse_indices_widgets))
        self.nse_all_widget.pack(anchor='w', padx=5, pady=2)
        # self.nse_all_widget.pack(anchor='w', padx=5, pady=2)
        self.nse_indices_vars = [tk.BooleanVar() for _ in range(len(nse_indices)+1)]
        
        for i, text in enumerate(nse_indices):
            var = self.nse_indices_vars[i]
            widget = ttk.Checkbutton(self.nse_indices_inner_frame, text=text, variable=var)
            widget.pack(anchor='w', padx=5, pady=2)
            self.nse_indices_widgets.append(widget)

        # Add widgets to the time period frame
        ttk.Label(time_period_frame, text="Start Date:").pack(anchor='w', padx=5, pady=2)
        self.start_date_entry = DateEntry(time_period_frame, width=12, background='darkblue', foreground='white', borderwidth=2)
        self.start_date_entry.pack(anchor='w', padx=5, pady=2)

        ttk.Label(time_period_frame, text="End Date:").pack(anchor='w', padx=5, pady=2)
        self.end_date_entry = DateEntry(time_period_frame, width=12, background='darkblue', foreground='white', borderwidth=2)
        self.end_date_entry.pack(anchor='w', padx=5, pady=2)

        ttk.Label(time_period_frame, text="Interval:").pack(anchor='w', padx=5, pady=2)
        self.interval_entry = ttk.Entry(time_period_frame)
        self.interval_entry.pack(anchor='w', padx=5, pady=2)

        self.next_button = ttk.Button(time_period_frame, text="Next >", command=self.calculations)
        self.next_button.pack(anchor='w', padx=5, pady=2)



    def calculations(self):
        global INTERVAL, STOCK_EXCHANGE, GROUP
        if STOCK_EXCHANGE == "bse":
            STOCK_INDEXES = [self.bse_indices_widgets[i].cget("text") for i, var in enumerate(self.bse_indices_vars) if var.get()]
            GROUP=[self.group_widgets[i].cget("text") for i,var in enumerate(self.group_vars) if var.get()]
        elif STOCK_EXCHANGE == "nse":
            STOCK_INDEXES = [self.nse_indices_widgets[i].cget("text") for i, var in enumerate(self.nse_indices_vars) if var.get()]
        else:
            STOCK_INDEXES = []
        self.START_DATE = self.start_date_entry.get_date()
        self.END_DATE = self.end_date_entry.get_date()
        try:
            interval_input = self.interval_entry.get().strip()
            if interval_input:
                INTERVAL = int(interval_input)
            else:
                INTERVAL = 0
        except ValueError:
            messagebox.showerror("Invalid Input", "Please enter a valid integer for the interval.")
            return
        
        self.START_DATE = pd.to_datetime(self.START_DATE)
        
        day = self.START_DATE.day_name()
        if(day == "Sunday"):
            self.START_DATE = self.START_DATE - pd.Timedelta(days=2)
        else:
            self.START_DATE = pd.to_datetime(self.START_DATE) - pd.Timedelta(days=1)
        self.START_DATE = self.START_DATE.date()
        print("stock_index............",STOCK_INDEXES[0])
        if STOCK_EXCHANGE == "nse": 
            if self.nse_all_var.get():
                query = f"SELECT * FROM mstr_nse_eod WHERE Date BETWEEN '{self.START_DATE}' AND '{self.END_DATE}'"
                final_df = pd.read_sql_query(query,self.conn)
                final_df = final_df[["NSETICKER","Date","Close"]]
                final_df["Date"] = pd.to_datetime(final_df["Date"])
            else:
                stock_df_list = []
                for stock_index in STOCK_INDEXES:
                    stock_index_name = INDEX_TO_TABLES_MAPPING[stock_index]
                    query = f"SELECT m.{STOCK_EXCHANGE.upper()}TICKER,il.Index_Code,e.Date,e.Close FROM  mstr_{STOCK_EXCHANGE}_stocks_list m INNER JOIN {stock_index_name} i on m.NSETICKER = i.NSETICKER INNER JOIN mstr_{STOCK_EXCHANGE}_index_list il ON i.Index_code = il.Index_code INNER JOIN mstr_{STOCK_EXCHANGE}_eod e ON m.NSETICKER = e.NSETICKER WHERE e.Date BETWEEN '{self.START_DATE}' AND '{self.END_DATE}'"
                    print("query ",query) #changed here
                    print("\n")
                    temp_df = pd.read_sql_query(query,self.conn)
                    stock_df_list.append(temp_df)

                final_df = pd.concat(stock_df_list,ignore_index=True)
                final_df["Date"] = pd.to_datetime(final_df["Date"])

        elif STOCK_EXCHANGE == "bse":
            if self.bse_all_var.get():
                query = f"SELECT * FROM mstr_bse_eod WHERE Date BETWEEN '{self.START_DATE}' AND '{self.END_DATE}'"
                final_df = pd.read_sql_query(query,self.conn)
                final_df = final_df[["BSETICKER","Date","Close"]]
                final_df["Date"] = pd.to_datetime(final_df["Date"])
            else: 
                stock_df_list = []
                for stock_index in STOCK_INDEXES:
                    stock_index_name = INDEX_TO_TABLES_MAPPING[stock_index]
                    for group in GROUP:
                        query = f"SELECT m.{STOCK_EXCHANGE.upper()}TICKER, il.Index_Code, e.Date, e.Close FROM mstr_{STOCK_EXCHANGE}_stocks_list m INNER JOIN {stock_index_name} i ON m.BSETICKER = i.BSETICKER INNER JOIN mstr_{STOCK_EXCHANGE}_index_list  il ON i.Index_Code = il.Index_Code  INNER JOIN mstr_{STOCK_EXCHANGE}_eod e on m.BSETICKER = e.BSETICKER WHERE e.Date BETWEEN '{self.START_DATE}' AND '{self.END_DATE}' AND m.BSE_Group LIKE'%{group}%'"
                        print(query)
                        temp_df = pd.read_sql_query(query,self.conn)
                        stock_df_list.append(temp_df)

                final_df = pd.concat(stock_df_list,ignore_index=True)
                final_df["Date"] = pd.to_datetime(final_df["Date"])

        print("final df",final_df)
        query2=f"select * from {self.DATABASE_NAME}.mstr_{STOCK_EXCHANGE}_indexes_eod WHERE Index_name = '{INDEX_TO_REGRESS_WITH}' AND Date BETWEEN '{self.START_DATE}' AND '{self.END_DATE}' "
        print("query2===",query2)
        temp = pd.read_sql_query(query2,self.conn)
        temp = temp[["Index_name","Date","Close"]]
        #renaming the column close to close market as to not confuse it with close of stocks 
        temp.rename(columns={
            "Close":"close_market"
        },inplace=True)
        temp["Date"] = pd.to_datetime(temp["Date"])
        print("temp length",len(temp))
        df313 = pd.merge(final_df,temp, how="inner",left_on="Date",right_on="Date" )
        df313 = pd.merge(final_df,temp, how="inner",left_on="Date",right_on="Date")
        df313 = df313[[f"{STOCK_EXCHANGE.upper()}TICKER","Date","Close","Index_name","close_market"]]
        df313.rename(columns={
            "Index_name":"index_to_regress_with",
            "Close": "Close_stock",
            "close_market": "Close_market_index"
            
        },inplace=True)
        print("len of df_512",len(df313))
        self.df_512 = {name:group for name,group in df313.groupby(f"{STOCK_EXCHANGE.upper()}TICKER")}


        for key in self.df_512:
            self.df_512[key] = self.df_512[key].iloc[::INTERVAL+1] # logic to implement interval 
            self.df_512[key] = self.df_512[key].sort_values(by="Date")
            self.df_512[key]["Previous_Close_stock"] = self.df_512[key]["Close_stock"].shift(1)
            self.df_512[key]["Previous_Close_market"] = self.df_512[key]["Close_market_index"].shift(1)
            self.df_512[key].dropna(inplace=True)

            self.df_512[key]["Return_Stock"] = ( self.df_512[key]["Close_stock"] - self.df_512[key]["Previous_Close_stock"] ) / self.df_512[key]["Previous_Close_stock"]
            self.df_512[key]["Return_market"] = ( self.df_512[key]["Close_market_index"] - self.df_512[key]["Previous_Close_market"] ) / self.df_512[key]["Previous_Close_market"]


            self.df_512[key]["returns_percentage"] = self.df_512[key]["Return_Stock"] * 100
            self.df_512[key]["returns_market_percentage"] = self.df_512[key]["Return_market"] * 100
            
        results = []
        for stock in self.df_512:

            if(len(self.df_512[stock]) < 5):  #added this to ensure stock has atleast 5 rows
                continue
            stock_return = self.df_512[stock]["returns_percentage"].to_numpy().reshape(-1,1)
            market_return = self.df_512[stock]["returns_market_percentage"].to_numpy().reshape(-1,1)

            model = LinearRegression()
            model.fit(market_return,stock_return)
            BETA = model.coef_[0][0]
            ALPHA = model.intercept_[0]
            mean_returns_stock = self.df_512[stock]["Return_Stock"].mean()
            std_deviation_stock = self.df_512[stock]["Return_Stock"].std()
            variance_stock = std_deviation_stock ** 2
            std_deviation_market = self.df_512[stock]["Return_market"].std()
            variance_market = std_deviation_market ** 2
            mean_returns_percentage_stock = self.df_512[stock]["returns_percentage"].mean()
            mean_returns_market = self.df_512[stock]["Return_market"].mean()
            mean_returns_market_percentage = mean_returns_market*100
            unsystematic_risk = variance_stock - ((BETA)**2  * variance_market )
            

            co_relation = self.df_512[stock]["Close_stock"].corr(self.df_512[stock]["Close_market_index"])
            total_risk = self.df_512[stock]["returns_percentage"].std()
            unsystematicrisk = np.sqrt((total_risk**2) - ((BETA)**2 * (self.df_512[stock]["returns_market_percentage"].std()**2)))
            systematic_risk = total_risk - unsystematicrisk            
            length_of_df = len(self.df_512[stock])

            results.append(
                
                { f"{STOCK_EXCHANGE.upper()}TICKER": self.df_512[stock][f"{STOCK_EXCHANGE.upper()}TICKER"].iloc[0],
                "no_of_rows": len(self.df_512[stock]),
                "mean_returns_stock": mean_returns_stock,
                "returns_percentage": mean_returns_percentage_stock,
                "annualized_returns_percentage (mean_daily_returns*365)": mean_returns_percentage_stock* 365,
                "Alpha": ALPHA,
                "BETA": BETA,
                "unsystematic_risk": unsystematic_risk,
                "total risk of stock(based on returns percentage)": total_risk, #added this here 
                "unsystematic risk (based on returns percentage)" : unsystematicrisk ,
                "systematic risk (based on returns percentage)" : systematic_risk,
                "std_deviation_stock": std_deviation_stock,
                "variance_stock": variance_stock ,
                "std_deviation_market": std_deviation_market,
                "variance_market" : variance_market,
                "co-relation (R^2)": co_relation,
                "mean_returns_market" : mean_returns_market,
                "mean_returns_market_percentage" : mean_returns_market_percentage
                }
            )
            
        print("len of df_512",len(self.df_512))
        table_1 = pd.DataFrame(results)
        self.table_1 = pd.DataFrame(results)
        self.window1_table=copy.deepcopy(self.table_1)
        print(self.window1_table)
        self.risk_free_return()

        
    def update_indices_state(self, *args):
        global STOCK_EXCHANGE, INDEX_TO_REGRESS_WITH, STOCK_INDEXES
        STOCK_EXCHANGE = self.exchange_var.get()
        
        if STOCK_EXCHANGE == "bse":
            INDEX_TO_REGRESS_WITH = "SENSEX"
            for widget in self.nse_indices_widgets:
                widget.config(state="disabled")
            for widget in self.bse_indices_widgets:
                widget.config(state="normal")
            for widget in self.group_widgets:
                widget.config(state="normal")
            self.bse_all_widget.config(state="normal")
            self.nse_all_widget.config(state="disabled")
        elif STOCK_EXCHANGE == "nse":
            INDEX_TO_REGRESS_WITH = "NIFTY_50"
            for widget in self.bse_indices_widgets:
                widget.config(state="disabled")
            for widget in self.group_widgets:
                widget.config(state="disabled")
            for widget in self.nse_indices_widgets:
                widget.config(state="normal")
            self.bse_all_widget.config(state="disabled")
            self.nse_all_widget.config(state="normal")
        else:
            for widget in self.bse_indices_widgets:
                widget.config(state="normal")
            for widget in self.nse_indices_widgets:
                widget.config(state="normal")
            for widget in self.group_widgets:
                widget.config(state="normal")
            self.bse_all_widget.config(state="normal")
            self.nse_all_widget.config(state="normal")
                
    def handle_all_indices(self, all_var, index_vars, index_widgets):
        all_selected = all_var.get()
        for var, widget in zip(index_vars, index_widgets):
            var.set(all_selected)
            widget.config(state="disabled" if all_selected else "normal")
  
    def risk_free_return(self):
        new_window=tk.Toplevel(self)
        new_window.title("initial_inputs")
        new_window.geometry("1200x800")
        
        global RATE_RISK_FREE_RETURN
        RATE_RISK_FREE_RETURN=0
        main_frame = ttk.Frame(new_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        input_frame = ttk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=(0, 20))
        
        # Risk Free Rate
        ttk.Label(input_frame, text="Risk Free Rate Of Return (%) [Ex.7.5]").grid(row=0, column=0, padx=10, pady=10, sticky="e")
        self.risk_free_entry = ttk.Entry(input_frame, width=20)
        self.risk_free_entry.grid(row=0, column=1, padx=10, pady=10)
        
        # Ticker    
        ttk.Label(input_frame, text="Ticker").grid(row=2, column=0, padx=10, pady=10, sticky="e")
        self.ticker_entry = ttk.Entry(input_frame, width=20)
        self.ticker_entry.grid(row=2, column=1, padx=10, pady=10)
        
        self.ticker_var = tk.StringVar()
        self.ticker_entry.config(textvariable=self.ticker_var)
        self.ticker_var.trace_add("write", self.on_ticker_entry_change)
        
        self.status_label = ttk.Label(input_frame, text="")
        self.status_label.grid(row=2, column=2, padx=10, pady=10, sticky="w")
        # Buttons frame
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=(0, 20))
        
        self.delete_button = ttk.Button(button_frame, text="Delete", command=self.delete_fun, style='my.TButton')
        self.delete_button.pack(side=tk.LEFT, padx=(10, 5))
        
        self.next_button = ttk.Button(button_frame, text="Next >", command=lambda: self.process_inputs(new_window), style='my.TButton')
        self.next_button.pack(side=tk.LEFT, padx=5)
        
        self.export_excel_button0 = ttk.Button(button_frame, text="Export Table 0 to Exel", command=lambda: self.export_to_excel(self.window1_table, "Table 0"), style='my.TButton')
        self.export_excel_button0.pack(side=tk.LEFT, padx=5)
        
        self.export_word_button0 = ttk.Button(button_frame, text="Export Table 0 to word", 
                                            command=lambda: self.export_to_word(self.window1_table, "Table 0"), 
                                            style='my.TButton')
        self.export_word_button0.pack(side=tk.LEFT, padx=5)
                
        ttk.Label(input_frame, text="Index Type (Stock Exchange):").grid(row=3, column=0, padx=10, pady=10, sticky="e")
        ttk.Label(input_frame, text=STOCK_EXCHANGE, font=("Times New Roman", 12, 'bold')).grid(row=3, column=1, padx=10, pady=10, sticky="w")
                # Table frame
        table_frame = ttk.Frame(main_frame)
        table_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create Treeview with scrollbar
        self.tree = ttk.Treeview(table_frame, style='Custom.Treeview')
        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        # Display the data
        self.display_table(self.window1_table)
        
        # Style configuration
        style = ttk.Style()
        style.configure('my.TButton')
        style.configure('Custom.Treeview', rowheight=25)
        style.configure('Custom.Treeview.Heading')
    def on_ticker_entry_change(self, *args):
        ticker = self.ticker_entry.get().strip().upper()
        # Clear previous selections
        for item in self.tree.get_children():
            self.tree.set(item, "Select", "")
        
        if ticker:
            found = False
            for item in self.tree.get_children():
                if self.tree.set(item, f"{STOCK_EXCHANGE.upper()}TICKER") == ticker:
                    self.tree.set(item, "Select", "[\u2713]")
                    self.tree.see(item)
                    found = True
                    break
            
            if not found:
                # Ticker not found in the table
                self.ticker_entry.config(background="light pink")  # Visual feedback
                self.ticker_entry.after(2000, lambda: self.ticker_entry.config(background="white"))  # Reset after 2 seconds
                # Optionally, show a message
                self.show_ticker_status("Ticker not found")
            else:
                self.ticker_entry.config(background="light green")  # Visual feedback for found ticker
                self.ticker_entry.after(2000, lambda: self.ticker_entry.config(background="white"))  # Reset after 2 seconds
                self.show_ticker_status("Ticker found")
        else:
            # Clear status if entry is empty
            self.show_ticker_status("")

    def show_ticker_status(self, message):
        # Assuming you have a label for status messages
        if hasattr(self, 'status_label'):
            self.status_label.config(text=message)
        else:
            print(message)  # Fallback if there's no status label 
    def process_inputs(self, window):
        global RATE_RISK_FREE_RETURN
        try:
            RATE_RISK_FREE_RETURN = float(self.risk_free_entry.get())
        except ValueError:
            print("INVALID INPUT")
            RATE_RISK_FREE_RETURN = 0
        
        RATE_RISK_FREE_RETURN = RATE_RISK_FREE_RETURN / 365
        print("Rate of Risk-Free Return", RATE_RISK_FREE_RETURN)
        
        self.table_1 = self.table_1[(self.table_1["returns_percentage"] > RATE_RISK_FREE_RETURN) & (self.table_1["BETA"] > 0)]
        self.table_1 = self.table_1.sort_values(by="returns_percentage", ascending=False)
        
        # Adding column excessive returns per beta
        self.table_1["excess returns"] = self.table_1["returns_percentage"] - RATE_RISK_FREE_RETURN
        self.table_1["excess returns to beta"] = (self.table_1["returns_percentage"] - RATE_RISK_FREE_RETURN) / self.table_1["BETA"]
        
        # Display the below dataframe in window_2_table_1
        window_2_table_1 = self.table_1[[f"{STOCK_EXCHANGE.upper()}TICKER", "returns_percentage", "excess returns", "BETA", "unsystematic_risk", "excess returns to beta"]].copy(deep=True)
        
        cutoff_table = self.table_1[[f"{STOCK_EXCHANGE.upper()}TICKER", "excess returns to beta", "BETA", "unsystematic_risk", "returns_percentage", "variance_market"]].sort_values(by="excess returns to beta", ascending=False)
        cutoff_table["(BETA^2)/(UNSYSTEMATIC_RISK)"] = (cutoff_table["BETA"]**2) / (cutoff_table["unsystematic_risk"])
        cutoff_table["cumulative_intermediate"] = cutoff_table["(BETA^2)/(UNSYSTEMATIC_RISK)"].cumsum()
        cutoff_table["treynor_ratio_times_A"] = cutoff_table["excess returns to beta"] * cutoff_table["(BETA^2)/(UNSYSTEMATIC_RISK)"]
        cutoff_table["cumulative_treynor"] = cutoff_table["treynor_ratio_times_A"].cumsum()
        cutoff_table["cutoff_point"] = (cutoff_table["variance_market"] * cutoff_table["cumulative_treynor"]) / (1 + (cutoff_table["variance_market"] * cutoff_table["cumulative_intermediate"]))
        
        # Display the table below in window_2_table_2
        window_2_table_2 = cutoff_table.copy(deep=True)
        
        cutoff_max = cutoff_table["cutoff_point"].max()
        next_table = cutoff_table[cutoff_table["excess returns to beta"] >= cutoff_max]
        next_table = next_table[[f"{STOCK_EXCHANGE.upper()}TICKER", "returns_percentage", "excess returns to beta", "BETA", "unsystematic_risk", "cutoff_point"]]
        next_table["cutoff_max"] = cutoff_max
        
        weighted_table = next_table.copy(deep=True)
        weighted_table["excess_trenor"] = weighted_table["excess returns to beta"] - weighted_table["cutoff_max"]
        weighted_table["Z_i"] = weighted_table["excess_trenor"] * (weighted_table["BETA"] / weighted_table["unsystematic_risk"])
        
        # Display this table in window_2_table_3
        window_2_table_3 = weighted_table[[f"{STOCK_EXCHANGE.upper()}TICKER", "cutoff_point", "cutoff_max", "Z_i"]].copy(deep=True)
        
        weighted_table["weight"] = weighted_table["Z_i"] / weighted_table["Z_i"].sum()
        weighted_table["weight_percentage"] = weighted_table["weight"] * 100
        weighted_table["weighted_return"] = weighted_table["returns_percentage"] * weighted_table["weight_percentage"]
        
        # Display the below table in window_2_table_4
        window_2_table_4 = weighted_table[[f"{STOCK_EXCHANGE.upper()}TICKER", "returns_percentage", "weight_percentage", "weighted_return"]].copy(deep=True)
        
        weighted_table.sort_values(by="weight", ascending=False)
        TOTAL_RETURN_ON_PORTFOLIO = weighted_table["weighted_return"].sum()
        print("total return on portoflio...",TOTAL_RETURN_ON_PORTFOLIO)
        self.window_2_table_1 = window_2_table_1
        self.window_2_table_2 = window_2_table_2
        self.window_2_table_3 = window_2_table_3
        self.window_2_table_4 = window_2_table_4 
        self.grid_view()

    def display_table(self, df):
        # Clear the existing data in the treeview
        for i in self.tree.get_children():
            self.tree.delete(i)
            
        if df.empty:
            print("DataFrame is empty.")
            return
        # Set up columns
        self.tree["columns"] =  ["Select"]+list(df.columns)
        self.tree["show"] = "headings"
        # Set column headings
        self.tree.heading("Select", text="Select")
        self.tree.column("Select", width=50, anchor="center")
        for column in df.columns:
            self.tree.heading(column, text=column, command=lambda c=column: self.sort_treeview(self.tree, c, False))
            self.tree.column(column, width=100)  # Adjust the width as needed

        # Add data to the treeview
        for index, row in df.iterrows():
            item = self.tree.insert("", "end", values=["[ ]"] + list(row))
            
        # Bind the checkbutton to update the tree
        self.tree.bind("<ButtonRelease-1>", self.on_click_select)
        
    def export_to_excel(self, df, table_name):
        # Ask user for save location
        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx")
        if file_path:
            # Create a Pandas Excel writer using XlsxWriter as the engine
            with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
                df.to_excel(writer, sheet_name=table_name, index=False)
                
                # Get the xlsxwriter workbook and worksheet objects
                workbook = writer.book
                worksheet = writer.sheets[table_name]
                
                # Add a header format
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BC',
                    'border': 1})

                # Write the column headers with the defined format
                for col_num, value in enumerate(df.columns.values):
                    worksheet.write(0, col_num, value, header_format)
                    
                # Auto-adjust columns' width
                for i, col in enumerate(df.columns):
                    column_len = max(df[col].astype(str).map(len).max(), len(col))
                    worksheet.set_column(i, i, column_len)

        print(f"{table_name} exported to {file_path}")
        messagebox.showinfo("Export Successful", f"{table_name} exported to {file_path}")

    def export_to_word(self, df, table_name):
            # Ask user for save location
            file_path = filedialog.asksaveasfilename(defaultextension=".docx")
            if file_path:
                # Create a new Word document
                doc = Document()
                
                # Add a title
                doc.add_heading(table_name, 0)
                
                # Add a table
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Table Grid'
                
                # Add the header row
                hdr_cells = table.rows[0].cells
                for i, column in enumerate(df.columns):
                    hdr_cells[i].text = column
                
                # Add the data rows
                for row in df.itertuples(index=False):
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value)
                
                # Save the document
                doc.save(file_path)
                
                print(f"{table_name} exported to {file_path}")
                messagebox.showinfo("Export Successful", f"{table_name} exported to {file_path}")
    def on_click_select(self, event):
        region = self.tree.identify("region", event.x, event.y)
        if region == "cell":
            column = self.tree.identify_column(event.x)
            if column == "#1":  # The Select column
                item = self.tree.identify_row(event.y)
                current = self.tree.set(item, "Select")
                new_value = "[\u2713]" if current == "[ ]" else "[ ]"
                self.tree.set(item, "Select", new_value)
    def sort_treeview(self, tree, col, reverse):
        data = [(tree.set(child, col), child) for child in tree.get_children('')]
        data.sort(reverse=reverse)
        for index, (val, child) in enumerate(data):
            tree.move(child, '', index)
        tree.heading(col, command=lambda: self.sort_treeview(tree, col, not reverse))
          
    def grid_view(self):
        new_window=tk.Toplevel(self)
        new_window.title("initial_inputs")
        new_window.geometry("1050x1050")
        print("table1")
        print(self.window_2_table_1)
        print("table2")
        print(self.window_2_table_2)
        print("table3")
        print(self.window_2_table_3)
        print("table4")
        print(self.window_2_table_4)
        # Create four frames for each table
        frames = [ttk.Frame(new_window) for _ in range(4)]
        for i, frame in enumerate(frames):
            row = i // 2
            col = i % 2
            frame.grid(row=row, column=col, padx=10, pady=10, sticky="nsew")
            new_window.grid_rowconfigure(row, weight=1)
            new_window.grid_columnconfigure(col, weight=1)

        # Create and populate treeviews for each table
        tables = [self.window_2_table_1, self.window_2_table_2, self.window_2_table_3, self.window_2_table_4]
        titles = ["Table 1", "Table 2", "Table 3", "Table 4"]
        
        for frame, table, title in zip(frames, tables, titles):
        # Create a label for the table title
            ttk.Label(frame, text=title).grid(row=0, column=0, pady=5)
            
            tree_frame = ttk.Frame(frame)
            tree_frame.grid(row=1, column=0, sticky="nsew")
            frame.grid_rowconfigure(1, weight=1)
            frame.grid_columnconfigure(0, weight=1)
            # Create Treeview
            tree = ttk.Treeview(tree_frame)
            tree.grid(row=0, column=0, sticky="nsew")
            tree_frame.grid_rowconfigure(0, weight=1)
            tree_frame.grid_columnconfigure(0, weight=1)

            # Add vertical scrollbar
            vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
            vsb.grid(row=0,column=1,sticky='ns')
            tree.configure(yscrollcommand=vsb.set)

            # Add horizontal scrollbar
            hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
            hsb.grid(row=1,column=0,sticky='ew')
            tree.configure(xscrollcommand=hsb.set)

            # Set up columns
            tree["columns"] = list(table.columns)
            tree["show"] = "headings"

            # Set column headings
            for column in table.columns:
                tree.heading(column, text=column, command=lambda c=column, t=tree: self.sort_treeview(t, c, False))
                tree.column(column, width=100)  # Adjust the width as needed

            # Add data to the treeview
            for _, row in table.iterrows():
                tree.insert("", "end", values=list(row))
                
            
            frame.pack_propagate(False)
            
            # Add a button at the bottom center
            button_frame = ttk.Frame(new_window)
            button_frame.grid(row=2, column=0, columnspan=2, pady=20)
    
            
            self.export_excel_button1 = ttk.Button(button_frame, text="Export Table 1 to Excel", 
                                            command=lambda: self.export_to_excel(self.window_2_table_1, "Table 1"), 
                                            style='my.TButton')
            self.export_excel_button1.pack(side=tk.LEFT, padx=5)
            self.export_word_button1 = ttk.Button(button_frame, text="Export Table 1 to word", 
                                     command=lambda: self.export_to_word(self.window_2_table_1, "Table 1"), 
                                                style='my.TButton')
            self.export_word_button1.pack(side=tk.LEFT, padx=5)
            self.export_excel_button2 = ttk.Button(button_frame, text="Export Table 4 to Excel", 
                                            command=lambda: self.export_to_excel(self.window_2_table_4, "Table 4"), 
                                            style='my.TButton')
            self.export_excel_button2.pack(side=tk.LEFT, padx=5)
            self.export_word_button2 = ttk.Button(button_frame, text="Export to Table 4 to word", 
                                                command=lambda: self.export_to_word(self.window_2_table_4, "Table 4"), 
                                                style='my.TButton')
            self.export_word_button2.pack(side=tk.LEFT, padx=5)
                        
            compare_button = ttk.Button(button_frame, text="Analyze Portfolio", command=self.comparison,style='my.TButton')
            compare_button.pack()
        
    def delete_fun(self):
        # Get selected items with the checkmark
        selected_items = [item for item in self.tree.get_children() if self.tree.set(item, "Select") == "[\u2713]"]
        
        # Check if there are no selected items
        if not selected_items:
            messagebox.showinfo("No Selection", "Please select rows to delete.")
            return
        
        # Ask for confirmation before deletion
        confirm = messagebox.askyesno("Confirm Deletion", "Are you sure you want to delete the selected rows?")
        if confirm:
            # Get the BSETICKER values for the selected items
            selected_nsetickers = [self.tree.set(item, f"{STOCK_EXCHANGE.upper()}TICKER") for item in selected_items]
            
            # Remove the selected tickers from table_1
            self.table_1 = self.table_1[~self.table_1[f"{STOCK_EXCHANGE.upper()}TICKER"].isin(selected_nsetickers)]
            
            # Update window_1_table with a deep copy of the modified table_1
            self.window1_table = copy.deepcopy(self.table_1)
            
            # Remove the selected items from the tree view
            for item in selected_items:
                self.tree.delete(item)
            
            # Show a success message
            messagebox.showinfo("Deletion Successful", f"{len(selected_items)} row(s) deleted.")
            self.ticker_var.set("")

    def comparison(self,new_start_date=None):
        self.portfolio_stocks_and_weights = {row[f"{STOCK_EXCHANGE.upper()}TICKER"]: row["weight_percentage"] for index, row in self.window_2_table_4.iterrows()}
        
        # Adjust START_DATE
        START_DATE = new_start_date if new_start_date else self.START_DATE
        START_DATE = (START_DATE + timedelta(days=1)).strftime("%Y-%m-%d")
        
        def get_stock_value(ticker: str, date: str, db_cursor=self.cursor, max_days: int = 365):
            DATE = datetime.strptime(date, "%Y-%m-%d").date()
            earliest_date = DATE - timedelta(days=max_days)
            current_date = DATE
            
            while current_date >= earliest_date:
                date_str = current_date.strftime("%Y-%m-%d")
                query = f"SELECT Close FROM mstr_{STOCK_EXCHANGE}_eod WHERE Date = '{date_str}' AND {STOCK_EXCHANGE.upper()}TICKER='{ticker}'"
                db_cursor.execute(query)
                result = db_cursor.fetchone()
                
                if result:
                    return result[0]
                else:
                    current_date -= timedelta(days=1)
            
            return None
        
        def make_portfolio(portfolio_with_weights: dict, start_date: str = START_DATE, db_cursor=self.cursor, portfolio_amount: int = 100000):
            portfolio = dict()
            for stock in portfolio_with_weights:
                amount_invested_in_stock = portfolio_amount * (portfolio_with_weights[stock]/100)
                stock_value_on_start_date = get_stock_value(stock, start_date, db_cursor=db_cursor)
                if stock_value_on_start_date is None:
                    print(stock, start_date, stock_value_on_start_date)
                    continue
                no_of_shares_of_stock = int(amount_invested_in_stock / stock_value_on_start_date)
                portfolio[stock] = no_of_shares_of_stock
            return portfolio
        
        def get_portfolio_value(date: str, portfolio_dict, db_cursor=self.cursor):
            values = []
            for stock in portfolio_dict:
                stock_value = get_stock_value(date=date, ticker=stock, db_cursor=db_cursor)
                total_value_of_stock = portfolio_dict[stock] * stock_value
                values.append(total_value_of_stock)
            value = np.array(values)
            return value.sum()
        
        def get_index_value(date:str, index_name: str = INDEX_TO_REGRESS_WITH, db_cursor=self.cursor, max_days:int = 365):
            DATE = datetime.strptime(date, "%Y-%m-%d").date()
            earliest_date = DATE - timedelta(days=max_days)
            current_date = DATE
            
            while current_date >= earliest_date:
                date_str = current_date.strftime("%Y-%m-%d")
                query = f"SELECT Close FROM mstr_{STOCK_EXCHANGE}_indexes_eod WHERE Date = '{date_str}' AND Index_name='{index_name}'"
                db_cursor.execute(query)
                result = db_cursor.fetchone()
                
                if result:
                    return result[0]
                else:
                    current_date -= timedelta(days=1)
            
            return None
        
        def calculate_portfolio_returns(start_date:str = START_DATE, end_date: str = self.END_DATE.strftime("%Y-%m-%d"), db_cursor=self.cursor, portfolio_dict=None):
            value_on_starting = get_portfolio_value(start_date, db_cursor=db_cursor, portfolio_dict=portfolio_dict)
            value_on_ending = get_portfolio_value(end_date, db_cursor=db_cursor, portfolio_dict=portfolio_dict)
            returns_percentage = ((value_on_ending-value_on_starting)/(value_on_starting))*100
            return returns_percentage
        
        def calculate_index_returns(start_date:str = START_DATE, end_date: str = self.END_DATE.strftime("%Y-%m-%d"), db_cursor=self.cursor, index_name: str = INDEX_TO_REGRESS_WITH):
            value_on_starting = get_index_value(start_date, db_cursor=db_cursor)
            value_on_ending = get_index_value(end_date, db_cursor=db_cursor)
            returns_percentage = ((value_on_ending-value_on_starting)/(value_on_starting))*100
            return returns_percentage
        
        def get_date_list(start_date:str = START_DATE, end_date:str = self.END_DATE.strftime("%Y-%m-%d"), date_format:str = "%Y-%m-%d"):
            start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
            end_date = datetime.strptime(end_date, "%Y-%m-%d").date()
            return [(start_date + timedelta(days=x)).strftime(date_format) for x in range((end_date - start_date).days + 1)]
        
        def get_data(start_date:str = START_DATE, end_date: str = self.END_DATE.strftime("%Y-%m-%d"), portfolio_dict=None, db_cursor=self.cursor):
            dates = get_date_list(start_date, end_date)
            datewise_returns = []
            datewise_index_returns = []
            for date in dates:
                print(f"got data for index for date: {date}")
                datewise_returns.append(calculate_portfolio_returns(end_date=date, db_cursor=db_cursor, portfolio_dict=portfolio_dict))
                datewise_index_returns.append(calculate_index_returns(end_date=date, db_cursor=db_cursor))
            dates = pd.date_range(start=start_date, end=end_date, freq="D")
            datewise_returns = np.array(datewise_returns)
            datewise_index_returns = np.array(datewise_index_returns)
            return (dates, datewise_returns, datewise_index_returns)
        
        def get_sharpe_ratio(risk_free_rate=RATE_RISK_FREE_RETURN):
            dates = get_date_list()
            temp_date = datetime.strptime(dates[0],"%Y-%m-%d") - timedelta(days=1)
            temp_date = temp_date.strftime("%Y-%m-%d")
            dates.insert(0,temp_date)
            daily_close = [get_portfolio_value(date,portfolio_dict=portfolio_dict) for date in dates]

            df = pd.DataFrame({'Date':dates,'Close':daily_close})
            df['Date'] = pd.to_datetime(df['Date'])
            df['previous_close'] = df['Close'].shift(1)
            df.dropna(inplace=True)
            df["daily_returns_percentage"] = ((df["Close"] - df["previous_close"]) / (df["previous_close"])) * 100 
            mean_daily_returns_percentage = df["daily_returns_percentage"].mean() * 242
            std_dev_daily_returns_percetage = df["daily_returns_percentage"].std() * np.sqrt(242)
            risk_free_rate = risk_free_rate * 365
            print("mean daily returns percentage ",mean_daily_returns_percentage) 
            print("risk free rate returns percentage",risk_free_rate)
            print("std dev of daily returns percentage" ,std_dev_daily_returns_percetage)
            sharpe_ratio = (mean_daily_returns_percentage - risk_free_rate) / (std_dev_daily_returns_percetage)
            return sharpe_ratio
        
        def get_treynor_ratio_table(portfolio_stocks_and_weights:dict = self.portfolio_stocks_and_weights, table_1 = self.table_1):
            tickers = pd.Series(portfolio_stocks_and_weights.keys())
            treynor_table = table_1[table_1[f"{STOCK_EXCHANGE.upper()}TICKER"].isin(tickers)][[f"{STOCK_EXCHANGE.upper()}TICKER","excess returns to beta","BETA",]]
            treynor_table.rename(columns={
                "excess returns to beta": "treynor ratio"
            },inplace=True)
            weight_list = []
            for ticker in treynor_table[f"{STOCK_EXCHANGE.upper()}TICKER"]:
                weight_list.append(portfolio_stocks_and_weights[ticker])
            treynor_table["weight_percentange"] = weight_list

            return treynor_table
        
        def calculate_and_get_portfolio_treynor_ratio(risk_free_rate = RATE_RISK_FREE_RETURN,portfolio_stocks_and_weights = self.portfolio_stocks_and_weights, start_date:str = self.START_DATE, end_date: str = self.END_DATE):
            table_1 = create_table_1_for_portfolio(portfolio_stocks_and_weights, start_date,end_date)
            treynor_table = get_treynor_ratio_table()
            treynor_table["weight"] = treynor_table["weight_percentange"] / 100
            treynor_table["weight*beta"] = treynor_table["weight"] * treynor_table["BETA"]
            portfolio_beta = treynor_table["weight*beta"].sum()
            mean_returs_portfolio_percentage = table_1["mean_returns_percentage_portfolio"].mean()
            risk_free_rate = risk_free_rate
            treynor_ratio = (mean_returs_portfolio_percentage - risk_free_rate ) / portfolio_beta
            return treynor_ratio
        
        def create_table_1_for_portfolio(portfolio_stocks_and_weights = self.portfolio_stocks_and_weights,start_date:str = START_DATE, end_date: str = self.END_DATE.strftime("%Y-%m-%d")):
            stocks = portfolio_stocks_and_weights.keys()
            date_list = get_date_list()
            data = []
            for date in date_list:
                portfolio_value = get_portfolio_value(date,portfolio_dict=portfolio_dict)
                index_value = get_index_value(date)
                data.append([date,portfolio_value,index_value])
            df = pd.DataFrame(data=data, columns=["Date","portfolio_value","index_value"])
            df["Date"] = pd.to_datetime(df["Date"])
            df["previous_day_value"] = df["portfolio_value"].shift(1)
            df.dropna(inplace=True)
            df["mean_returns_percentage_portfolio"] = ((df["portfolio_value"] - df["previous_day_value"])/(df["previous_day_value"])) * 100   
            df = df[df["mean_returns_percentage_portfolio"] != 0.000000]
            return df
        
        portfolio_dict = make_portfolio(self.portfolio_stocks_and_weights, START_DATE, self.cursor)
        self.dates, self.datewise_returns, self.datewise_index_returns = get_data(START_DATE, self.END_DATE.strftime("%Y-%m-%d"), portfolio_dict, self.cursor)
        self.sharpe_ratio=get_sharpe_ratio()
        self.treynor_ratio=calculate_and_get_portfolio_treynor_ratio()
        print("DATES.............",self.dates)
        print("DATEWISE_RETURNS.............",self.datewise_returns)
        print("DATEWISE_INDEXES..................",self.datewise_index_returns)
        print("sharpe_ratio=",self.sharpe_ratio)
        print("treynors ratio....",self.treynor_ratio)
        self.initial_graph_setup(self.dates, self.datewise_returns, self.datewise_index_returns, self.sharpe_ratio, self.treynor_ratio)

        
        # Display the plot in a new Tkinter window
        
    def plot_results(self,ticker, range:tuple = (-10,10), precision:int = 100, no_of_bins_hist:int = 25):
            self.fig, self.axes = plt.subplots(nrows=1, ncols=2, figsize=(10,6))
            self.fig.suptitle(f"{ticker}", fontsize=16)
            print(self.df_512)
            print(self.df_512.keys())
            print(self.df_512[ticker])
            # print(self.df_512[ticker])
            n, bins, patches = self.axes[0].hist(self.df_512[ticker]["returns_percentage"], bins=no_of_bins_hist, edgecolor="black", alpha=0.7)
            self.axes[0].hist(self.df_512[ticker]["returns_percentage"], bins=no_of_bins_hist, edgecolor="black")
            self.axes[0].set_xlabel("returns percentage")
            self.axes[0].set_ylabel("frequency")
            
            mu = self.df_512[ticker]["returns_percentage"].mean()  # Mean
            sigma = self.df_512[ticker]["returns_percentage"].std()  # Standard deviation
            x = np.linspace(range[0], range[1], precision)
            
            pdf = norm.pdf(x, mu, sigma)
            
            self.axes[1].plot(x, pdf)
            self.axes[1].set_xlabel("returns_percentage")
            self.axes[1].set_ylabel("probability density")
            
            for ax in self.axes:
                ax.grid(True, linestyle='--', alpha=0.7)
                ax.tick_params(axis='both', which='major', labelsize=10)
                
            plt.tight_layout()
            plt.show()
        
    def display_analysis(self):
            analysis_window = tk.Toplevel(self)
            analysis_window.title("Stock Analysis")

            # Create dropdown for stock selection
            stock_list = list(self.portfolio_stocks_and_weights.keys())
            selected_stock = tk.StringVar(analysis_window)
            selected_stock.set(stock_list[0])  # Set default value
            stock_dropdown = ttk.Combobox(analysis_window, textvariable=selected_stock, values=stock_list)
            stock_dropdown.pack(pady=10)
            
            # Create input fields for other parameters
            min_range = tk.Entry(analysis_window)
            min_range.insert(0, "-10")
            min_range.pack(pady=5)
            tk.Label(analysis_window, text="Min Range").pack()

            max_range = tk.Entry(analysis_window)
            max_range.insert(0, "10")
            max_range.pack(pady=5)
            tk.Label(analysis_window, text="Max Range").pack()

            precision = tk.Entry(analysis_window)
            precision.insert(0, "100")
            precision.pack(pady=5)
            tk.Label(analysis_window, text="Precision").pack()

            no_of_bins = tk.Entry(analysis_window)
            no_of_bins.insert(0, "25")
            no_of_bins.pack(pady=5)
            tk.Label(analysis_window, text="Number of Histogram Bins").pack()

            def run_analysis():
                ticker = selected_stock.get() #chanqge for bse float(int()) and nse as it is
                if(STOCK_EXCHANGE=='bse'):
                    ticker=int(float(ticker))
                range_tuple = (float(min_range.get()), float(max_range.get()))
                precision_value = int(precision.get())
                bins = int(no_of_bins.get())
                self.plot_results(ticker, range=range_tuple, precision=precision_value, no_of_bins_hist=bins)

            analyze_button = ttk.Button(analysis_window, text="Analyze", command=run_analysis)
            analyze_button.pack(pady=20)

    
    def on_mouse_move(self,event):
        if event.inaxes:
            x, y = event.xdata, event.ydata
            date = mdates.num2date(x).strftime('%Y-%m-%d')
            self.hline.set_ydata([y,y])
            self.vline.set_xdata([x,x])
            self.text_box.set_text(f'Date: {date}\nReturns: {y:.2f}%')
            self.remove_annotation()
            self.fig.canvas.draw_idle()
            
    def on_click(self, event):
        global current_annotation
        # Check if event has 'inaxes' attribute (only true for Matplotlib events)
        if not hasattr(event, 'inaxes') or event.inaxes is None:
            return
        
        self.remove_annotation()
        date = mdates.num2date(event.xdata).strftime('%Y-%m-%d')
        returns = event.ydata
        current_annotation = self.ax.annotate(
            f'Date: {date}\nReturns: {returns:.2f}%', 
            xy=(event.xdata, event.ydata),
            xytext=(10, 10), textcoords='offset points',
            bbox=dict(boxstyle='round,pad=0.5', fc='yellow', alpha=0.5),
            arrowprops=dict(arrowstyle='->', connectionstyle='arc3,rad=0')
        )
        self.fig.canvas.draw_idle()
    def remove_annotation(self):
        global current_annotation
        if current_annotation:
            current_annotation.remove()
            current_annotation = None
            
    def initial_graph_setup(self, dates, datewise_returns, datewise_index_returns, sharpe_ratio, treynor_ratio):
        # Create a new window
        new_window = tk.Toplevel(self)
        new_window.title("Portfolio Comparison")
        new_window.geometry('1000x600')
        new_window.grab_set()
        self.fig, self.ax = plt.subplots(figsize=(10, 6))
        # Create frames for layout
        left_frame = tk.Frame(new_window, width=700, height=500)
        right_frame = tk.Frame(new_window, width=300, height=500, padx=10)
        left_frame.pack(side="left", fill="both", expand=True)
        right_frame.pack(side="right", fill="y", padx=10, pady=10)

        # LEFT FRAME: Graph
        canvas = FigureCanvasTkAgg(self.fig, master=left_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

        # Plotting the graph
        self.ax.plot(dates, datewise_returns, label='Portfolio Returns')
        self.ax.plot(dates, datewise_index_returns, label='Index Returns')
        self.ax.set_xlabel('Date')
        self.ax.set_ylabel('Expected Returns (%)')
        self.ax.set_title('Portfolio Returns vs Index Returns')
        self.ax.legend()

        # Display ratios on the graph
        self.fig.text(0.5, 0.01, f"Annualized Sharpe Ratio = {sharpe_ratio:.4f}", ha='center', va='bottom', fontsize=12, fontweight='bold')
        self.fig.text(0.5, 0.04, f"Treynor's Ratio = {treynor_ratio:.4f}", ha='center', va='bottom', fontsize=12, fontweight='bold')
        plt.subplots_adjust(bottom=0.15)

        # Cursor and annotations
        self.curs = Cursor(self.ax, useblit=True, color='red', linewidth=1)
        self.hline = self.ax.axhline(y=0, color='k', lw=0.8, ls='--')
        self.vline = self.ax.axvline(x=dates[0], color='k', lw=0.8, ls='--')
        self.text_box = self.ax.text(0.02, 0.98, '', transform=self.ax.transAxes, va='top', ha='left', 
                        bbox=dict(facecolor='white', edgecolor='none', alpha=0.65))
        self.fig.canvas.mpl_connect('motion_notify_event', self.on_mouse_move)
        self.fig.canvas.mpl_connect('button_press_event', self.on_click)

        # RIGHT FRAME: Data input and buttons
        start_date_var = StringVar()
        Label(right_frame, text="Start Date (YYYY-MM-DD):").pack(pady=5)
        date_entry = DateEntry(right_frame, textvariable=start_date_var, date_pattern="yyyy-mm-dd")
        date_entry.pack(pady=10)
        
        # Buttons
        analyze_button_port = ttk.Button(right_frame, text="Analyze Portfolio", command=lambda: self.comparison(datetime.strptime(start_date_var.get(), "%Y-%m-%d")))
        analyze_button_port.pack(pady=5)

        analyze_button_stocks = ttk.Button(right_frame, text="Analyze Stocks", command=self.display_analysis)
        analyze_button_stocks.pack(pady=5)

        # Final adjustments
        left_frame.pack_propagate(False)  # Keep fixed size for frames
        right_frame.pack_propagate(False)
        self.fig.canvas.draw_idle()      
    def update_plot( self,dates, datewise_returns, datewise_index_returns, sharpe_ratio, treynor_ratio):
            new_window = tk.Toplevel(self)
            new_window.title("Portfolio Comparison")
            canvas = FigureCanvasTkAgg(self.fig, master=new_window)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both",expand=True)
            button_frame=tk.Frame(new_window)
            start_date_var = StringVar()
            Label(button_frame, text="Start Date (YYYY-MM-DD):").pack()
            DateEntry(button_frame, textvariable=start_date_var, date_pattern="yyyy-mm-dd").pack()
            start_date_str = start_date_var.get()
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
            button_frame.pack(pady=20)
            analyze_button_port = ttk.Button(button_frame, text="Analyze Portfolio ", command=self.comparison(start_date))
            analyze_button_port.pack()
            analyze_button = ttk.Button(button_frame, text="Analyze Stocks ", command=self.display_analysis)
            analyze_button.pack()
            self.ax.clear()
            self.ax.plot(dates, datewise_returns, label='Portfolio Returns')
            self.ax.plot(dates, datewise_index_returns, label='Index Returns')
            self.ax.set_xlabel('Date')
            self.ax.set_ylabel('Expected Returns (%)')
            self.ax.set_title('Portfolio Returns vs Index Returns')
            self.ax.legend()
            
            self.fig.texts.clear()  # Clear existing text annotations
            self.fig.text(0.5, 0.01, f"Annualized Sharpe Ratio = {sharpe_ratio:.4f}", ha='center', va='bottom', fontsize=12, fontweight='bold')
            self.fig.text(0.5, 0.04, f"Treynor's Ratio = {treynor_ratio:.4f}", ha='center', va='bottom', fontsize=12, fontweight='bold')
            
            plt.subplots_adjust(bottom=0.15)

            # Recreate cursor, hline, vline, and text_box
            self.curs = Cursor(self.ax, useblit=True, color='red', linewidth=1)
            self.hline = self.ax.axhline(y=0, color='k', lw=0.8, ls='--')
            self.vline = self.ax.axvline(x=dates[0], color='k', lw=0.8, ls='--')
            
            self.text_box = self.ax.text(0.02, 0.98, '', transform=self.ax.transAxes, va='top', ha='left', 
                        bbox=dict(facecolor='white', edgecolor='none', alpha=0.65))

            self.fig.canvas.draw()
            self.fig.canvas.mpl_connect('motion_notify_event', self.on_mouse_move)
            self.fig.canvas.mpl_connect('button_press_event', self.on_click)

